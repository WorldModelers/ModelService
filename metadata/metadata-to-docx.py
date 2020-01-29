from docx import Document
from docx.shared import Inches
import yaml
import json
import os
import glob

def gen_doc(model):
    model_name = model['label']
    
    document = Document()

    document.add_heading(f'{model_name}', 0)

    p = document.add_paragraph(model['description'].replace('\n','\n\n'))

    document.add_paragraph(
        f"Model ID: {model['id']}", style='List Bullet'
    )
    document.add_paragraph(
        f"Model Maintainer: {model['maintainer'].get('name','')}, {model['maintainer'].get('email','')}", style='List Bullet'       
    )
    document.add_paragraph(
        f"Model Category: {', '.join(model['category'])}", style='List Bullet'       
    )
    
    document.add_heading('Outputs', level=1)
    for o in model['outputs']:
        p = document.add_paragraph()
        p.add_run(f"{o['name']}: ").bold = True
        p.paragraph_format.space_after = 0
        document.add_paragraph(
            f"Description: {o['description']}", style='List Bullet'
        )                
        document.add_paragraph(
            f"Units: {o['units']}", style='List Bullet'
        )        
        
    document.add_heading('Parameters', level=1)        
    if 'parameters' in model:        
        for p_ in model['parameters']:
            p = document.add_paragraph()
            p.add_run(f"{p_['name']}: ").bold = True
            p.paragraph_format.space_after = 0        
            document.add_paragraph(
                f"Description: {p_['description']}", style='List Bullet'
            )
            document.add_paragraph(
                f"Type: {p_['metadata']['type']}", style='List Bullet'
            )   

            # choice parameter
            if p_['metadata']['type'] == 'ChoiceParameter':
                choices = [str(i) for i in p_['metadata']['choices']]
                document.add_paragraph(
                    f"Choices: {', '.join(choices)}", style='List Bullet'
                )

           # number parameter
            if p_['metadata']['type'] == 'NumberParameter':
                document.add_paragraph(
                    f"Min/Max: {p_['metadata'].get('minimum','None')}, {p_['metadata'].get('maximum','None')}", style='List Bullet'
                ) 

            document.add_paragraph(
                f"Default: {p_['metadata'].get('default','None')}", style='List Bullet'
            )          

    else: #model has no parameters
        p = document.add_paragraph("Model has no parameters.")

    document.save(f"metadata-docx/{model['id']}.docx")


if __name__ == "__main__":

    if not os.path.exists('metadata-docx'):
        os.mkdir('metadata-docx')

    models = glob.glob('models/*yaml')

    for m in models:

        with open(f'{m}', 'r') as stream:
            model = yaml.safe_load(stream)
            gen_doc(model)